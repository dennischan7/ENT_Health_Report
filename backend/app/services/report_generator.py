"""
Report Generator Service for creating professional DOCX reports.

This module provides the ReportGenerator class for generating comprehensive
enterprise health diagnosis reports in Microsoft Word format (.docx).

Features:
    - Title page with company information
    - Executive summary section
    - Financial metrics tables
    - Bar charts and line charts embedded from matplotlib
    - Peer comparison analysis
    - Risk assessment
    - Recommendations

Dependencies:
    - python-docx: For DOCX generation
    - matplotlib: For chart generation

Example:
    >>> from app.services.report_generator import ReportGenerator
    >>> from app.schemas.report import ReportData
    >>>
    >>> generator = ReportGenerator()
    >>> doc_path = generator.generate(report_data)
    >>> print(f"Report saved to: {doc_path}")
"""

import io
import logging
import os
import time
from datetime import date
from decimal import Decimal
from pathlib import Path
from typing import List, Optional, Tuple

import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings
from app.schemas.report import (
    ReportData,
    FinancialMetric,
    TrendData,
    ChartData,
    ExecutiveSummary,
    FinancialMetricsSection,
    PeerComparisonSection,
    RiskAssessmentSection,
    RecommendationsSection,
)
from app.services.report_styles import (
    Colors, Fonts, ChartSettings, PageSettings,
    set_table_header_style, set_table_alternating_rows,
    set_cell_value_format, highlight_target_row,
    format_number, format_change_rate,
    setup_page_margins, add_page_header, add_page_footer,
)


# Use non-interactive backend for matplotlib (required for server environments)
matplotlib.use("Agg")

# Set Chinese font support for matplotlib
# Use fonts that are installed in the Docker container
plt.rcParams["font.sans-serif"] = ["WenQuanYi Zen Hei", "Noto Sans CJK SC", "Noto Sans CJK JP", "SimHei", "Microsoft YaHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["font.family"] = "sans-serif"

logger = logging.getLogger(__name__)


class ReportGeneratorError(Exception):
    """Exception raised for report generation errors."""

    pass


class ReportGenerator:
    """
    Generator for enterprise health diagnosis reports in DOCX format.

    This class provides methods for creating professional reports with:
    - Title page
    - Executive summary
    - Financial metrics tables
    - Embedded charts (bar and line)
    - Peer comparison analysis
    - Risk assessment
    - Recommendations

    Attributes:
        reports_dir: Directory where generated reports are saved.
        chart_width: Default chart width in inches.
        chart_height: Default chart height in inches.

    Example:
        >>> generator = ReportGenerator(reports_dir="./reports")
        >>> doc_path = generator.generate(report_data)
    """

    def __init__(
        self,
        reports_dir: Optional[str] = None,
        chart_width: float = 6.0,
        chart_height: float = 4.0,
    ):
        """
        Initialize the report generator.

        Args:
            reports_dir: Directory to save generated reports.
                        Defaults to 'reports/' in the backend directory.
            chart_width: Default width for embedded charts in inches.
            chart_height: Default height for embedded charts in inches.
        """
        if reports_dir:
            self.reports_dir = Path(reports_dir)
        else:
            # Default to reports directory in backend
            self.reports_dir = Path(__file__).parent.parent.parent / "reports"

        # Ensure reports directory exists
        self.reports_dir.mkdir(parents=True, exist_ok=True)

        self.chart_width = chart_width
        self.chart_height = chart_height

        logger.info(f"ReportGenerator initialized with reports_dir: {self.reports_dir}")

    def generate(
        self,
        report_data: ReportData,
        output_filename: Optional[str] = None,
    ) -> str:
        """
        Generate a DOCX report from the provided data.

        Args:
            report_data: ReportData object containing all report sections.
            output_filename: Optional custom filename for the output file.
                           If not provided, a filename is generated automatically.

        Returns:
            str: Absolute path to the generated DOCX file.

        Raises:
            ReportGeneratorError: If report generation fails.
        """
        start_time = time.time()

        try:
            # Generate output filename if not provided
            if not output_filename:
                output_filename = self._generate_filename(report_data)

            output_path = self.reports_dir / output_filename

            logger.info(f"Starting report generation: {output_filename}")

            # Create document
            doc = Document()

            # Set up document styles
            self._setup_document_styles(doc)

            # Apply professional page margins
            setup_page_margins(doc)

            # 1. Title Page
            self._add_title_page(doc, report_data)

            # 2. Table of Contents placeholder
            self._add_table_of_contents(doc)

            # 3. Executive Summary
            self._add_executive_summary(doc, report_data.executive_summary)

            # 4. Financial Metrics Tables
            self._add_financial_metrics(doc, report_data.financial_metrics)

            # 5. Bar Charts
            if report_data.bar_charts:
                self._add_charts_section(
                    doc, report_data.bar_charts, "关键指标对比", chart_type="bar"
                )

            # 6. Line Charts (Trends)
            if report_data.trends:
                self._add_trend_charts(doc, report_data.trends)

            # 7. Peer Comparison Analysis
            if report_data.peer_comparison:
                self._add_peer_comparison(doc, report_data.peer_comparison)

            # 8. Risk Assessment
            self._add_risk_assessment(doc, report_data.risk_assessment)

            # 9. Recommendations
            self._add_recommendations(doc, report_data.recommendations)

            # 10. Footer with generation info
            self._add_footer(doc, report_data)

            # Save document
            doc.save(str(output_path))

            # Verify file was created
            if not output_path.exists():
                raise ReportGeneratorError(f"Failed to create report file: {output_path}")

            file_size = output_path.stat().st_size
            generation_time = time.time() - start_time

            logger.info(
                f"Report generated successfully: {output_path} "
                f"(size: {file_size} bytes, time: {generation_time:.2f}s)"
            )

            return str(output_path.absolute())

        except Exception as e:
            logger.error(f"Report generation failed: {e}", exc_info=True)
            raise ReportGeneratorError(f"Failed to generate report: {e}") from e

    def _generate_filename(self, report_data: ReportData) -> str:
        """Generate a unique filename for the report."""
        company_code = report_data.company_code
        report_date = report_data.report_date.strftime("%Y%m%d")
        report_type = report_data.report_type.value

        filename = f"{company_code}_{report_type}_{report_date}.docx"
        return filename

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up custom document styles."""
        styles = doc.styles

        # Title style
        if "ReportTitle" not in [s.name for s in styles]:
            title_style = styles.add_style("ReportTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "Microsoft YaHei"
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

        # Heading 1 style modification
        heading1 = styles["Heading 1"]
        heading1.font.name = "Microsoft YaHei"
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)

        # Heading 2 style modification
        heading2 = styles["Heading 2"]
        heading2.font.name = "Microsoft YaHei"
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(51, 51, 51)

        # Normal style modification
        normal = styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(11)

    def _add_title_page(self, doc: Document, report_data: ReportData) -> None:
        """Add a title page to the document."""
        # Add some spacing at the top
        for _ in range(3):
            doc.add_paragraph()

        # Main title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("企业健康度诊断报告")
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Company name
        doc.add_paragraph()
        company = doc.add_paragraph()
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company.add_run(report_data.company_name)
        run.font.size = Pt(24)
        run.font.bold = True

        # Company code
        code = doc.add_paragraph()
        code.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = code.add_run(f"股票代码: {report_data.company_code}")
        run.font.size = Pt(16)

        # Industry
        industry = doc.add_paragraph()
        industry.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = industry.add_run(f"所属行业: {report_data.industry_name}")
        run.font.size = Pt(14)

        # Add spacing
        for _ in range(4):
            doc.add_paragraph()

        # Report type
        report_type_names = {
            "full_diagnosis": "完整诊断报告",
            "financial_analysis": "财务分析报告",
            "peer_comparison": "同业对比报告",
            "risk_assessment": "风险评估报告",
        }
        type_name = report_type_names.get(
            report_data.report_type.value, report_data.report_type.value
        )
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = type_para.add_run(f"报告类型: {type_name}")
        run.font.size = Pt(14)

        # Report years
        if report_data.report_years:
            years = doc.add_paragraph()
            years.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = years.add_run(f"分析周期: {report_data.report_years}")
            run.font.size = Pt(14)

        # Report date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"报告日期: {report_data.report_date.strftime('%Y年%m月%d日')}")
        run.font.size = Pt(14)

        # Page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Add a table of contents placeholder."""
        doc.add_heading("目录", level=1)

        toc_items = [
            "一、执行摘要",
            "二、财务指标分析",
            "三、关键指标对比图表",
            "四、三年趋势分析",
            "五、同业对比分析",
            "六、风险评估",
            "七、建议与措施",
        ]

        for item in toc_items:
            para = doc.add_paragraph(item)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, summary: ExecutiveSummary) -> None:
        """Add executive summary section."""
        doc.add_heading("一、执行摘要", level=1)

        # Overall rating
        rating_para = doc.add_paragraph()
        rating_para.add_run("综合评级: ").bold = True
        rating_para.add_run(summary.overall_rating)

        if summary.overall_score:
            score_para = doc.add_paragraph()
            score_para.add_run("综合得分: ").bold = True
            score_para.add_run(f"{summary.overall_score}分")

        # Summary text
        doc.add_paragraph()
        summary_heading = doc.add_paragraph()
        summary_heading.add_run("摘要").bold = True
        doc.add_paragraph(summary.summary_text)

        # Key strengths
        if summary.key_strengths:
            doc.add_paragraph()
            strengths_heading = doc.add_paragraph()
            strengths_heading.add_run("主要优势").bold = True
            for strength in summary.key_strengths:
                para = doc.add_paragraph(strength, style="List Bullet")

        # Key risks
        if summary.key_risks:
            doc.add_paragraph()
            risks_heading = doc.add_paragraph()
            risks_heading.add_run("主要风险").bold = True
            for risk in summary.key_risks:
                para = doc.add_paragraph(risk, style="List Bullet")

        # Recommendation
        doc.add_paragraph()
        rec_heading = doc.add_paragraph()
        rec_heading.add_run("综合建议").bold = True
        doc.add_paragraph(summary.recommendation)

        doc.add_page_break()

    def _add_financial_metrics(self, doc: Document, metrics: FinancialMetricsSection) -> None:
        """Add financial metrics tables section."""
        doc.add_heading("二、财务指标分析", level=1)

        # Profitability metrics
        if metrics.profitability:
            doc.add_heading("2.1 盈利能力指标", level=2)
            self._add_metrics_table(doc, metrics.profitability)

        # Solvency metrics
        if metrics.solvency:
            doc.add_heading("2.2 偿债能力指标", level=2)
            self._add_metrics_table(doc, metrics.solvency)

        # Operation metrics
        if metrics.operation:
            doc.add_heading("2.3 运营能力指标", level=2)
            self._add_metrics_table(doc, metrics.operation)

        # Growth metrics
        if metrics.growth:
            doc.add_heading("2.4 成长能力指标", level=2)
            self._add_metrics_table(doc, metrics.growth)

        doc.add_page_break()

    def _add_metrics_table(self, doc: Document, metrics: List[FinancialMetric]) -> None:
        """Add a professionally styled table of financial metrics."""
        if not metrics:
            return

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        headers = ["指标名称", "当前值", "单位", "同比变化", "行业平均"]
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Apply professional header style
        set_table_header_style(table)

        # Data rows
        for metric in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric.name
            set_cell_value_format(row_cells[1], self._format_decimal(metric.value), is_number=True)
            row_cells[2].text = metric.unit

            # Format change rate with color
            change_str, is_positive = format_change_rate(metric.change_rate)
            set_cell_value_format(row_cells[3], change_str, is_number=True, is_positive=is_positive)

            set_cell_value_format(row_cells[4], self._format_decimal(metric.industry_avg), is_number=True)

        # Apply alternating row colors
        set_table_alternating_rows(table)

        doc.add_paragraph()  # Add spacing after table

    def _format_decimal(self, value: Optional[Decimal]) -> str:
        """Format a decimal value for display."""
        if value is None:
            return "-"
        # Format with 2 decimal places and thousand separators
        return f"{value:,.2f}"

    def _format_change_rate(self, value: Optional[Decimal]) -> str:
        """Format a change rate percentage."""
        if value is None:
            return "-"
        sign = "+" if value > 0 else ""
        return f"{sign}{value:.2f}%"

    def _add_charts_section(
        self,
        doc: Document,
        charts_data: List[ChartData],
        section_title: str,
        chart_type: str = "bar",
    ) -> None:
        """Add a section with embedded charts and analysis text."""
        doc.add_heading("三、关键指标对比图表", level=1)

        for chart_data in charts_data:
            # Add chart title
            doc.add_heading(chart_data.title, level=2)

            # Create and embed chart
            chart_image = self._create_chart(chart_data, chart_type)
            if chart_image:
                doc.add_picture(chart_image, width=Inches(self.chart_width))
                doc.add_paragraph()  # Spacing after chart

            # Add analysis text for the chart
            analysis_text = self._generate_chart_analysis(chart_data)
            if analysis_text:
                doc.add_paragraph(analysis_text)

        doc.add_page_break()

    def _generate_chart_analysis(self, chart_data: ChartData) -> str:
        """Generate analysis text for a chart based on its data."""
        try:
            datasets = chart_data.datasets
            if not datasets or len(datasets) < 2:
                return ""

            target_data = datasets[0].get("data", [])
            avg_data = datasets[1].get("data", [])
            labels = chart_data.x_labels

            if not target_data or not avg_data:
                return ""

            # Find metrics where target is significantly different from average
            above_avg = []
            below_avg = []
            close_avg = []

            for i, (label, target_val, avg_val) in enumerate(zip(labels, target_data, avg_data)):
                if target_val is None or avg_val is None or avg_val == 0:
                    continue

                diff_pct = (target_val / avg_val - 1) * 100

                if diff_pct > 20:
                    above_avg.append((label, diff_pct))
                elif diff_pct < -20:
                    below_avg.append((label, diff_pct))
                else:
                    close_avg.append(label)

            # Build analysis text
            parts = []

            if above_avg:
                metrics_str = "、".join([f"{m}（高出{p:.1f}%）" for m, p in above_avg[:3]])
                parts.append(f"**领先指标**：{metrics_str}等指标显著高于行业平均水平。")

            if below_avg:
                metrics_str = "、".join([f"{m}（低于{abs(p):.1f}%）" for m, p in below_avg[:3]])
                parts.append(f"**落后指标**：{metrics_str}等指标明显低于行业平均，需重点关注和改进。")

            if close_avg:
                metrics_str = "、".join(close_avg[:3])
                parts.append(f"**接近平均**：{metrics_str}等指标与行业平均水平接近，整体表现稳定。")

            return " ".join(parts) if parts else ""

        except Exception as e:
            logger.warning(f"Failed to generate chart analysis: {e}")
            return ""

    def _create_chart(self, chart_data: ChartData, chart_type: str = "bar") -> Optional[io.BytesIO]:
        """
        Create a chart image and return as BytesIO.

        Args:
            chart_data: Chart data including labels and datasets.
            chart_type: Type of chart ('bar', 'line', 'pie', 'radar').

        Returns:
            BytesIO containing the chart image, or None if creation fails.
        """
        try:
            fig, ax = plt.subplots(figsize=(self.chart_width, self.chart_height))

            x_labels = chart_data.x_labels
            x = range(len(x_labels))

            # Use professional color palette
            colors = ChartSettings.COLORS

            if chart_type == "bar":
                # Grouped bar chart for multiple datasets
                bar_width = 0.8 / len(chart_data.datasets)
                for i, dataset in enumerate(chart_data.datasets):
                    offset = (i - len(chart_data.datasets) / 2 + 0.5) * bar_width
                    values = [float(v) if v is not None else 0 for v in dataset.get("data", [])]
                    bars = ax.bar(
                        [xi + offset for xi in x],
                        values,
                        bar_width,
                        label=dataset.get("label", f"Series {i + 1}"),
                        color=colors[i % len(colors)],
                    )
                    # Add data labels on bars
                    for bar, val in zip(bars, values):
                        if val != 0:
                            height = bar.get_height()
                            ax.annotate(f'{val:,.0f}',
                                        xy=(bar.get_x() + bar.get_width() / 2, height),
                                        xytext=(0, 3),
                                        textcoords="offset points",
                                        ha='center', va='bottom',
                                        fontsize=8)

                ax.set_xticks(x)
                ax.set_xticklabels(x_labels, rotation=45, ha="right")

            elif chart_type == "line":
                for i, dataset in enumerate(chart_data.datasets):
                    values = [float(v) if v is not None else 0 for v in dataset.get("data", [])]
                    line = ax.plot(
                        x,
                        values,
                        marker="o",
                        label=dataset.get("label", f"Series {i + 1}"),
                        color=colors[i % len(colors)],
                        linewidth=2,
                        markersize=6,
                    )
                    # Add data labels on points
                    for xi, val in zip(x, values):
                        ax.annotate(f'{val:,.0f}',
                                    xy=(xi, val),
                                    xytext=(0, 5),
                                    textcoords="offset points",
                                    ha='center', va='bottom',
                                    fontsize=8)

                ax.set_xticks(x)
                ax.set_xticklabels(x_labels)

            elif chart_type == "radar":
                # Radar chart for multi-dimensional comparison
                return self._create_radar_chart(chart_data)

            ax.set_title(chart_data.title, fontsize=12, fontweight="bold", color=Colors.PRIMARY_HEX)
            ax.legend(loc="best", fontsize=9)
            ax.grid(True, alpha=0.3)

            # Format y-axis with thousand separators
            ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: format(int(x), ",")))

            plt.tight_layout()

            # Save to BytesIO with high DPI
            image_stream = io.BytesIO()
            fig.savefig(image_stream, format="png", dpi=ChartSettings.DPI, bbox_inches="tight")
            image_stream.seek(0)
            plt.close(fig)

            return image_stream

        except Exception as e:
            logger.error(f"Failed to create chart '{chart_data.title}': {e}")
            plt.close("all")
            return None

    def _create_radar_chart(self, chart_data: ChartData) -> Optional[io.BytesIO]:
        """Create a radar chart for multi-dimensional comparison."""
        try:
            import numpy as np

            labels = chart_data.x_labels
            num_vars = len(labels)

            # Compute angle for each axis
            angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
            angles += angles[:1]  # Complete the loop

            fig, ax = plt.subplots(figsize=(self.chart_width, self.chart_height), subplot_kw=dict(polar=True))

            colors = ChartSettings.COLORS

            for i, dataset in enumerate(chart_data.datasets):
                values = [float(v) if v is not None else 0 for v in dataset.get("data", [])]
                values += values[:1]  # Complete the loop

                ax.plot(angles, values, 'o-', linewidth=2,
                        label=dataset.get("label", f"Series {i + 1}"),
                        color=colors[i % len(colors)])
                ax.fill(angles, values, alpha=0.25, color=colors[i % len(colors)])

            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(labels, fontsize=9)
            ax.set_title(chart_data.title, fontsize=12, fontweight="bold",
                         color=Colors.PRIMARY_HEX, pad=20)
            ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.0), fontsize=9)

            plt.tight_layout()

            image_stream = io.BytesIO()
            fig.savefig(image_stream, format="png", dpi=ChartSettings.DPI, bbox_inches="tight")
            image_stream.seek(0)
            plt.close(fig)

            return image_stream

        except Exception as e:
            logger.error(f"Failed to create radar chart '{chart_data.title}': {e}")
            plt.close("all")
            return None

    def _add_trend_charts(self, doc: Document, trends: List[TrendData]) -> None:
        """Add trend line charts section with analysis text."""
        doc.add_heading("四、三年趋势分析", level=1)

        for trend in trends:
            if not trend.data:
                continue

            # Add chart title
            doc.add_heading(trend.metric_name, level=2)

            # Create trend chart
            chart_image = self._create_trend_chart(trend)
            if chart_image:
                doc.add_picture(chart_image, width=Inches(self.chart_width))
                doc.add_paragraph()

            # Add analysis text for the trend
            analysis_text = self._generate_trend_analysis(trend)
            if analysis_text:
                doc.add_paragraph(analysis_text)

        doc.add_page_break()

    def _generate_trend_analysis(self, trend: TrendData) -> str:
        """Generate analysis text for a trend chart."""
        try:
            data = trend.data
            if len(data) < 2:
                return f"{trend.metric_name}数据不足，无法进行趋势分析。"

            # Calculate year-over-year changes
            changes = []
            for i in range(1, len(data)):
                prev_val = float(data[i-1].value) if data[i-1].value else 0
                curr_val = float(data[i].value) if data[i].value else 0

                if prev_val != 0:
                    change_pct = (curr_val / prev_val - 1) * 100
                    changes.append((data[i].year, change_pct, curr_val, prev_val))

            if not changes:
                return ""

            # Determine overall trend
            positive_changes = [c for c in changes if c[1] > 0]
            negative_changes = [c for c in changes if c[1] < 0]

            latest_year, latest_change, latest_val, _ = changes[-1]
            first_val = float(data[0].value) if data[0].value else 0
            last_val = float(data[-1].value) if data[-1].value else 0

            if first_val != 0:
                total_change = (last_val / first_val - 1) * 100
            else:
                total_change = 0

            # Build analysis text
            parts = []

            # Overall trend
            if total_change > 10:
                parts.append(f"**整体趋势**：{trend.metric_name}呈上升趋势，{data[0].year}年至{data[-1].year}年累计增长{total_change:.1f}%。")
            elif total_change < -10:
                parts.append(f"**整体趋势**：{trend.metric_name}呈下降趋势，{data[0].year}年至{data[-1].year}年累计下降{abs(total_change):.1f}%。")
            else:
                parts.append(f"**整体趋势**：{trend.metric_name}整体保持稳定，变动幅度在{abs(total_change):.1f}%以内。")

            # Year-by-year analysis
            if len(changes) > 0:
                year_changes = []
                for year, change_pct, curr_val, prev_val in changes:
                    direction = "增长" if change_pct > 0 else "下降"
                    year_changes.append(f"{year-1}至{year}年{direction}{abs(change_pct):.1f}%")
                parts.append(f"**年度变化**：{'；'.join(year_changes)}。")

            # Latest value
            parts.append(f"**最新数据**：{data[-1].year}年{trend.metric_name}为{last_val:,.2f}{trend.unit}。")

            return " ".join(parts)

        except Exception as e:
            logger.warning(f"Failed to generate trend analysis: {e}")
            return ""

    def _create_trend_chart(self, trend: TrendData) -> Optional[io.BytesIO]:
        """Create a trend line chart from TrendData with professional styling."""
        try:
            fig, ax = plt.subplots(figsize=(self.chart_width, self.chart_height))

            years = [d.year for d in trend.data]
            values = [float(d.value) if d.value is not None else 0 for d in trend.data]

            ax.plot(years, values, marker="o", linewidth=2, color=Colors.PRIMARY_HEX,
                    markersize=8, markerfacecolor=Colors.PRIMARY_LIGHT_HEX)

            # Add data labels
            for year, val in zip(years, values):
                ax.annotate(f'{val:,.0f}',
                            xy=(year, val),
                            xytext=(0, 8),
                            textcoords="offset points",
                            ha='center', va='bottom',
                            fontsize=9)

            ax.set_title(f"{trend.metric_name}趋势", fontsize=12, fontweight="bold",
                         color=Colors.PRIMARY_HEX)
            ax.set_xlabel("年份", fontsize=10)
            ax.set_ylabel(f"{trend.metric_name} ({trend.unit})", fontsize=10)
            ax.grid(True, alpha=0.3)

            # Format y-axis with thousand separators
            ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: format(int(x), ",")))

            plt.tight_layout()

            # Save to BytesIO with high DPI
            image_stream = io.BytesIO()
            fig.savefig(image_stream, format="png", dpi=ChartSettings.DPI, bbox_inches="tight")
            image_stream.seek(0)
            plt.close(fig)

            return image_stream

        except Exception as e:
            logger.error(f"Failed to create trend chart for {trend.metric_name}: {e}")
            plt.close("all")
            return None

    def _add_peer_comparison(self, doc: Document, peer: PeerComparisonSection) -> None:
        """Add peer comparison analysis section with professional styling."""
        doc.add_heading("五、同业对比分析", level=1)

        # Peer companies list
        if peer.peer_companies:
            doc.add_heading("5.1 对比公司", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            header_cells[0].text = "公司代码"
            header_cells[1].text = "公司名称"
            header_cells[2].text = "所属行业"

            # Apply professional header style
            set_table_header_style(table)

            for company in peer.peer_companies:
                row = table.add_row().cells
                row[0].text = company.company_code
                row[1].text = company.company_name
                row[2].text = company.industry_name

            # Apply alternating row colors
            set_table_alternating_rows(table)

            doc.add_paragraph()

        # Comparison metrics with enhanced styling
        if peer.comparison_metrics:
            doc.add_heading("5.2 对比指标", level=2)
            self._add_metrics_table(doc, peer.comparison_metrics)

            # Add radar chart for peer comparison if we have metrics
            if len(peer.comparison_metrics) >= 3:
                self._add_peer_radar_chart(doc, peer)

        # Analysis text
        if peer.analysis_text:
            doc.add_heading("5.3 分析结论", level=2)
            # Split analysis into paragraphs if it contains newlines
            paragraphs = peer.analysis_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        # Industry ranking with visual indicator
        if peer.ranking_in_industry:
            rank_para = doc.add_paragraph()
            rank_para.add_run("行业排名: ").bold = True
            rank_run = rank_para.add_run(f"第{peer.ranking_in_industry}名")
            # Color code based on ranking
            if peer.ranking_in_industry <= 3:
                rank_run.font.color.rgb = Colors.SUCCESS
            elif peer.ranking_in_industry <= 10:
                rank_run.font.color.rgb = Colors.WARNING
            else:
                rank_run.font.color.rgb = Colors.DANGER

        doc.add_page_break()

    def _add_peer_radar_chart(self, doc: Document, peer: PeerComparisonSection) -> None:
        """Add radar chart for peer comparison metrics."""
        # Create radar chart data from comparison metrics
        labels = []
        target_values = []
        industry_values = []

        for metric in peer.comparison_metrics[:8]:  # Limit to 8 metrics for readability
            if metric.value is not None:
                labels.append(metric.name[:6])  # Truncate long names
                # Normalize to percentage of industry average
                if metric.industry_avg and metric.industry_avg != 0:
                    target_values.append(float(metric.value) / float(metric.industry_avg) * 100)
                else:
                    target_values.append(100)
                industry_values.append(100)  # Industry average is baseline

        if len(labels) >= 3:
            radar_data = ChartData(
                chart_type="radar",
                title="核心指标对比雷达图",
                x_labels=labels,
                datasets=[
                    {"label": "目标企业", "data": target_values},
                    {"label": "行业平均", "data": industry_values},
                ]
            )

            chart_image = self._create_radar_chart(radar_data)
            if chart_image:
                doc.add_picture(chart_image, width=Inches(self.chart_width))
                doc.add_paragraph()

    def _add_risk_assessment(self, doc: Document, risk: RiskAssessmentSection) -> None:
        """Add risk assessment section."""
        doc.add_heading("六、风险评估", level=1)

        # Overall risk level
        level_para = doc.add_paragraph()
        level_para.add_run("整体风险等级: ").bold = True
        risk_level = risk.overall_risk_level
        level_para.add_run(risk_level)

        # Color code the risk level
        if risk_level == "低":
            level_para.runs[-1].font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif risk_level == "中":
            level_para.runs[-1].font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            level_para.runs[-1].font.color.rgb = RGBColor(255, 0, 0)  # Red

        # Financial risk
        doc.add_heading("6.1 财务风险", level=2)
        doc.add_paragraph(risk.financial_risk)

        # Operational risk
        doc.add_heading("6.2 经营风险", level=2)
        doc.add_paragraph(risk.operational_risk)

        # Market risk
        doc.add_heading("6.3 市场风险", level=2)
        doc.add_paragraph(risk.market_risk)

        # Risk warnings
        if risk.risk_warnings:
            doc.add_heading("6.4 风险预警事项", level=2)
            for warning in risk.risk_warnings:
                para = doc.add_paragraph(warning, style="List Bullet")

        doc.add_page_break()

    def _add_recommendations(self, doc: Document, recs: RecommendationsSection) -> None:
        """Add recommendations section."""
        doc.add_heading("七、建议与措施", level=1)

        # Strategic recommendations
        if recs.strategic_recommendations:
            doc.add_heading("7.1 战略建议", level=2)
            for rec in recs.strategic_recommendations:
                doc.add_paragraph(rec, style="List Bullet")

        # Financial recommendations
        if recs.financial_recommendations:
            doc.add_heading("7.2 财务建议", level=2)
            for rec in recs.financial_recommendations:
                doc.add_paragraph(rec, style="List Bullet")

        # Operational recommendations
        if recs.operational_recommendations:
            doc.add_heading("7.3 经营建议", level=2)
            for rec in recs.operational_recommendations:
                doc.add_paragraph(rec, style="List Bullet")

        # Action items
        if recs.action_items:
            doc.add_heading("7.4 行动事项", level=2)
            for i, item in enumerate(recs.action_items, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. {item}")

    def _add_footer(self, doc: Document, report_data: ReportData) -> None:
        """Add footer with generation information."""
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)

        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(
            f"本报告由企业健康度智能诊断平台自动生成\n"
            f"生成时间: {date.today().strftime('%Y年%m月%d日')}\n"
            f"报告版本: v1.0"
        ).font.size = Pt(9)

    def get_report_path(self, filename: str) -> Path:
        """Get the full path for a report file."""
        return self.reports_dir / filename

    def cleanup_old_reports(self, days_old: int = 30) -> int:
        """
        Remove reports older than specified days.

        Args:
            days_old: Number of days after which reports should be deleted.

        Returns:
            int: Number of files deleted.
        """
        import time

        deleted_count = 0
        cutoff_time = time.time() - (days_old * 24 * 60 * 60)

        for file_path in self.reports_dir.glob("*.docx"):
            if file_path.stat().st_mtime < cutoff_time:
                try:
                    file_path.unlink()
                    deleted_count += 1
                    logger.info(f"Deleted old report: {file_path}")
                except Exception as e:
                    logger.warning(f"Failed to delete {file_path}: {e}")

        return deleted_count


# Convenience function
def get_report_generator(
    reports_dir: Optional[str] = None,
    chart_width: float = 6.0,
    chart_height: float = 4.0,
) -> ReportGenerator:
    """
    Get a ReportGenerator instance with specified settings.

    Args:
        reports_dir: Directory to save generated reports.
        chart_width: Default width for embedded charts in inches.
        chart_height: Default height for embedded charts in inches.

    Returns:
        ReportGenerator: A configured report generator instance.
    """
    return ReportGenerator(
        reports_dir=reports_dir,
        chart_width=chart_width,
        chart_height=chart_height,
    )
