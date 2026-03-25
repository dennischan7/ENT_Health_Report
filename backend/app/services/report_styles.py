"""
Professional Document Styles for Report Generation.

This module provides professional styling constants and helper functions
for generating high-quality DOCX reports.

Features:
    - Professional color palette
    - Table styling helpers
    - Font and paragraph formatting
    - Chart color schemes
"""

from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from typing import Optional


# =============================================================================
# Color Palette
# =============================================================================

class Colors:
    """Professional color palette for enterprise reports."""

    # Primary colors (for DOCX)
    PRIMARY = RGBColor(0, 51, 102)          # Deep blue #003366
    PRIMARY_LIGHT = RGBColor(0, 112, 192)   # Light blue #0070C0

    # Semantic colors (for DOCX)
    SUCCESS = RGBColor(0, 128, 0)           # Green #008000
    WARNING = RGBColor(255, 165, 0)         # Orange #FFA500
    DANGER = RGBColor(192, 0, 0)            # Red #C00000

    # Neutral colors (for DOCX)
    BLACK = RGBColor(0, 0, 0)
    DARK_GRAY = RGBColor(64, 64, 64)        # #404040
    GRAY = RGBColor(128, 128, 128)          # #808080
    LIGHT_GRAY = RGBColor(242, 242, 242)    # #F2F2F2
    WHITE = RGBColor(255, 255, 255)

    # Table colors (for DOCX)
    TABLE_HEADER_BG = RGBColor(0, 51, 102)  # Deep blue
    TABLE_HEADER_TEXT = RGBColor(255, 255, 255)  # White
    TABLE_ROW_ALT = RGBColor(242, 242, 242)  # Light gray
    TABLE_BORDER = RGBColor(166, 166, 166)  # Medium gray

    # Chart colors for matplotlib (hex strings)
    PRIMARY_HEX = "#003366"          # Deep blue
    PRIMARY_LIGHT_HEX = "#0070C0"    # Light blue
    SUCCESS_HEX = "#008000"          # Green
    WARNING_HEX = "#FFA500"          # Orange
    DANGER_HEX = "#C00000"           # Red

    # Chart colors (professional palette for matplotlib)
    CHART_COLORS = [
        "#003366",  # Deep blue
        "#0070C0",  # Light blue
        "#00B050",  # Green
        "#FFC000",  # Gold
        "#C00000",  # Red
        "#7030A0",  # Purple
        "#00B0F0",  # Cyan
        "#FF6600",  # Orange
    ]


# =============================================================================
# Font Settings
# =============================================================================

class Fonts:
    """Font settings for document elements."""

    # Font families
    TITLE_FONT = "Microsoft YaHei"
    BODY_FONT = "Microsoft YaHei"
    ENGLISH_FONT = "Arial"

    # Font sizes
    TITLE_SIZE = Pt(28)
    HEADING1_SIZE = Pt(16)
    HEADING2_SIZE = Pt(14)
    HEADING3_SIZE = Pt(12)
    BODY_SIZE = Pt(11)
    SMALL_SIZE = Pt(9)
    CAPTION_SIZE = Pt(8)

    # Line spacing
    SINGLE_LINE = 1.0
    ONE_POINT_FIVE_LINE = 1.5
    DOUBLE_LINE = 2.0


# =============================================================================
# Page Settings
# =============================================================================

class PageSettings:
    """Page layout settings."""

    # Margins (in cm)
    TOP_MARGIN = Cm(2.54)
    BOTTOM_MARGIN = Cm(2.54)
    LEFT_MARGIN = Cm(3.17)
    RIGHT_MARGIN = Cm(3.17)

    # Paper size (A4)
    PAGE_WIDTH = Cm(21.0)
    PAGE_HEIGHT = Cm(29.7)


# =============================================================================
# Chart Settings
# =============================================================================

class ChartSettings:
    """Chart generation settings."""

    DPI = 200
    WIDTH_INCHES = 6.0
    HEIGHT_INCHES = 4.0

    # Colors for matplotlib
    COLORS = Colors.CHART_COLORS


# =============================================================================
# Table Styling Helpers
# =============================================================================

def set_table_header_style(table, header_row_idx: int = 0) -> None:
    """
    Apply professional header style to table header row.

    Args:
        table: docx Table object
        header_row_idx: Index of header row (default: 0)
    """
    row = table.rows[header_row_idx]
    for cell in row.cells:
        # Set background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')
        cell._tc.get_or_add_tcPr().append(shading)

        # Set text formatting
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = Colors.WHITE
                run.font.size = Fonts.BODY_SIZE


def set_table_alternating_rows(table, start_row: int = 1) -> None:
    """
    Apply alternating row colors to table data rows.

    Args:
        table: docx Table object
        start_row: First data row index (default: 1, after header)
    """
    for i, row in enumerate(table.rows[start_row:], start=start_row):
        if (i - start_row) % 2 == 1:  # Odd rows
            for cell in row.cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F2F2F2')
                cell._tc.get_or_add_tcPr().append(shading)


def set_cell_value_format(cell, value: str, is_number: bool = False,
                          is_positive: Optional[bool] = None) -> None:
    """
    Format cell value with proper alignment and optional color.

    Args:
        cell: docx Cell object
        value: Text value to set
        is_number: Whether this is a numeric value (right-align if True)
        is_positive: For change rates, color green if True, red if False
    """
    cell.text = ""
    para = cell.paragraphs[0]

    if is_number:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run(value)
    run.font.size = Fonts.BODY_SIZE

    if is_positive is True:
        run.font.color.rgb = Colors.SUCCESS
    elif is_positive is False:
        run.font.color.rgb = Colors.DANGER


def highlight_target_row(table, row_idx: int) -> None:
    """
    Highlight the target enterprise row in comparison table.

    Args:
        table: docx Table object
        row_idx: Index of the row to highlight
    """
    row = table.rows[row_idx]
    for cell in row.cells:
        # Light blue background for target
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'DEEAF6')
        cell._tc.get_or_add_tcPr().append(shading)

        # Bold text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True


def format_number(value, decimal_places: int = 2, unit: str = "") -> str:
    """
    Format number with thousand separators and optional unit.

    Args:
        value: Numeric value (int, float, Decimal)
        decimal_places: Number of decimal places
        unit: Optional unit suffix

    Returns:
        Formatted string (e.g., "1,234,567.89元")
    """
    if value is None:
        return "-"

    try:
        if decimal_places > 0:
            formatted = f"{float(value):,.{decimal_places}f}"
        else:
            formatted = f"{int(value):,}"
    except (ValueError, TypeError):
        return str(value)

    if unit:
        formatted += unit

    return formatted


def format_change_rate(value, decimal_places: int = 2) -> str:
    """
    Format change rate percentage with sign and color indicator.

    Args:
        value: Change rate value
        decimal_places: Number of decimal places

    Returns:
        Tuple of (formatted_string, is_positive)
    """
    if value is None:
        return "-", None

    try:
        rate = float(value)
        sign = "+" if rate > 0 else ""
        formatted = f"{sign}{rate:.{decimal_places}f}%"
        return formatted, rate >= 0
    except (ValueError, TypeError):
        return str(value), None


# =============================================================================
# Document Setup Helpers
# =============================================================================

def setup_page_margins(doc) -> None:
    """
    Set professional page margins for the document.

    Args:
        doc: docx Document object
    """
    for section in doc.sections:
        section.top_margin = PageSettings.TOP_MARGIN
        section.bottom_margin = PageSettings.BOTTOM_MARGIN
        section.left_margin = PageSettings.LEFT_MARGIN
        section.right_margin = PageSettings.RIGHT_MARGIN


def add_page_header(doc, title: str, date_str: str) -> None:
    """
    Add header to document pages.

    Args:
        doc: docx Document object
        title: Report title
        date_str: Report date string
    """
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = header_para.add_run(f"{title} | {date_str}")
        run.font.size = Fonts.SMALL_SIZE
        run.font.color.rgb = Colors.GRAY


def add_page_footer(doc) -> None:
    """
    Add footer with page number and disclaimer.

    Args:
        doc: docx Document object
    """
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = footer_para.add_run("第 ")
        run.font.size = Fonts.SMALL_SIZE
        run.font.color.rgb = Colors.GRAY

        # Page number field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"
        run._r.append(instr_text)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)

        run2 = footer_para.add_run(" 页")
        run2.font.size = Fonts.SMALL_SIZE
        run2.font.color.rgb = Colors.GRAY